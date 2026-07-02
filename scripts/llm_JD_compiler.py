#!/usr/bin/env python3
"""
LLM-based Job Description compiler.

Reads a JD from a .docx or plain-text file and emits a structured JSON of the
hiring signals we need for downstream ranking.

The model is asked to emit YAML (which chat models generate more reliably than
strict JSON, especially with nested fields and long lists). We parse the YAML
and write the result to disk as JSON so downstream code stays JSON-native.

Extracted fields:

  - required_technical_skills : hard skills the JD explicitly demands
  - preferred_technical_skills: nice-to-haves / adjacent tech
  - min_years_experience      : lower bound of acceptable YoE (float)
  - max_years_experience      : upper bound if stated, else null
  - target_role_titles        : titles that fit the role
  - domain_keywords           : product/domain terms (RAG, ranking, embeddings, …)
  - disqualifiers             : explicit dealbreakers the JD calls out
  - location                  : required/preferred locations
  - seniority                 : junior | mid | senior | staff | principal
  - visa_sponsorship          : bool, defaults to false unless JD explicitly offers it
  - expected_projects         : concise, concrete artefacts the candidate is expected
                                to build IN THIS ROLE (not their past work). Each item
                                should name a specific system/component, not a category.
  - position_requirement_type : "product" if the JD clearly describes what will be built
                                / owned; "profile" if it only lists skills + experience
  - company_type              : startup | mid_stage | mnc | service | agency | other
  - exclude_companies         : list of companies to down-weight (not hard-filter).
                                Downstream blender assigns a negative score
                                proportional to months at any listed company.

The LLM is called via an OpenAI-compatible endpoint so it can point at
Fireworks (or OpenAI, vLLM, LM Studio, etc.) with only env vars changed.

Usage:
    export OPENAI_API_KEY=fw_...            # Fireworks API key
    export OPENAI_BASE_URL=https://api.fireworks.ai/inference/v1
    export JD_MODEL="accounts/fireworks/models/llama-v3p1-70b-instruct"

    python scripts/llm_JD_compiler.py \\
        --in India_runs_data_and_ai_challenge/job_description.docx \\
        --out scripts/jd_compiled.json

Environment env conda: py3.10
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path

import yaml


def _load_dotenv(path: Path) -> None:
    """Tiny .env loader — avoids a python-dotenv dependency.
    Only sets keys that are not already in the environment."""
    if not path.exists():
        return
    for raw in path.read_text().splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, _, val = line.partition("=")
        key = key.strip()
        val = val.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = val


# Load .env sitting next to this script (LLM_API_KEY / LLM_BASE_URL / LLM_MODEL_NAME).
_load_dotenv(Path(__file__).with_name(".env"))


def read_jd(path: Path) -> str:
    """Load JD text from .docx or plain text (.txt, .md, .rtf-as-text, stdin)."""
    if str(path) == "-":
        return sys.stdin.read()

    suffix = path.suffix.lower()
    if suffix == ".docx":
        try:
            import docx  # python-docx
        except ImportError as e:
            raise SystemExit(
                "python-docx not installed. Run: pip install python-docx"
            ) from e
        doc = docx.Document(str(path))
        parts: list[str] = []
        parts.extend(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(t for t in parts if t and t.strip())

    return path.read_text(encoding="utf-8")


PROMPT_PATH = Path(__file__).with_name("jd_compiler_prompt.md")


def _load_prompt_file(path: Path) -> tuple[str, str] | None:
    """Parse scripts/jd_compiler_prompt.md into (system_prompt, yaml_template).

    Sections are delimited by `## SYSTEM_PROMPT` and `## OUTPUT_YAML_TEMPLATE`
    (case-sensitive, at the start of a line). Text outside those sections is
    ignored so the file can carry human notes at the top.

    Returns None if the file is missing/malformed — caller falls back to the
    hardcoded strings below."""
    if not path.exists():
        return None
    try:
        text = path.read_text()
    except OSError:
        return None

    # Line-anchored search so mentions of the marker inside the file's HTML
    # header comment don't shadow the real section headings.
    sys_re = re.compile(r"^## SYSTEM_PROMPT[ \t]*$", re.MULTILINE)
    tpl_re = re.compile(r"^## OUTPUT_YAML_TEMPLATE[ \t]*$", re.MULTILINE)
    sys_match = sys_re.search(text)
    tpl_match = tpl_re.search(text)
    if not sys_match or not tpl_match or tpl_match.start() < sys_match.end():
        return None
    sys_body = text[sys_match.end() : tpl_match.start()].strip()
    tpl_body = text[tpl_match.end() :].strip()
    if not sys_body or not tpl_body:
        return None
    # Preserve a trailing newline on the template so YAML parses identically
    # to the pre-externalization behavior.
    return sys_body, tpl_body + "\n"


_FALLBACK_SYSTEM_PROMPT = """You are a hiring-signal extractor. Read a job description and produce a strict YAML document capturing what a ranking system needs to score candidates against this role.

Rules:
- Return ONLY valid YAML. No prose, no markdown fences, no ```yaml``` blocks.
- Use YAML block style (key: value on each line, list items as `  - item`). Do NOT use JSON-style flow syntax.
- If a field is not stated in the JD, use `null` (or `[]` for lists).
- Skills MUST be short canonical tokens (e.g. "PyTorch", "RAG", "vector search", "BM25", "embeddings", "fine-tuning LLMs"). Do NOT paraphrase full sentences into a skill.
- Split "required" vs "preferred" strictly:
  * required: the JD says this is needed / must-have / disqualifier if absent.
  * preferred: nice-to-have / bonus / "would be great".
- min_years_experience is a float (e.g. 5.0). If the JD gives a range like "5-9 years", set min to 5 and max to 9. If it says "senior" without numbers, infer a sensible min (>=5.0).
- target_role_titles: titles a matching candidate could plausibly currently hold.
- disqualifiers: extract EXPLICIT dealbreakers only (e.g. "no pure research background").
- Do not invent skills the JD doesn't mention.
- visa_sponsorship: set true ONLY if the JD explicitly says the employer sponsors work visas / work permits / relocation-from-abroad. Default to false when the JD is silent.
- expected_projects: 2-6 SHORT, CONCRETE artefacts the candidate is expected to BUILD OR OWN in this role. These items will be used AS QUERIES in a vector search against candidates' past career_history[].description text — so they must READ LIKE something an engineer would write about a past project they shipped. Optimize for semantic retrieval, not for how the JD phrased it.
  * Write each item as a NATURAL PROJECT DESCRIPTION: name the system, the technical approach, and the outcome. 10-25 words.
  * Use engineer-voice, past-tense-ready phrasing (e.g. "Built X using Y to do Z"). Do NOT start with verbs like "Ship", "Own", "Design" — those are JD-voice, not resume-voice.
  * Include the concrete tech stack + the problem domain in each item so dense embeddings have strong lexical + semantic anchors.
  * NO meta / process artefacts (do NOT emit "evaluation infrastructure", "A/B testing framework", "feedback loops", "team mentoring", "architecture design"). Those don't match career_history descriptions well. Only emit BUILT SYSTEMS.
  * GOOD:
      - "hybrid retrieval system combining BM25 and dense embeddings for candidate ranking"
      - "LLM-based re-ranker over top-k retrieval results with prompt-engineered relevance scoring"
      - "semantic search over resume corpus using sentence-transformer embeddings and vector database"
      - "learning-to-rank model trained on recruiter click and hire signals"
  * BAD (too broad, meta, or JD-voice — do NOT emit):
      - "production ranking systems at scale"
      - "evaluation infrastructure (offline benchmarks, online A/B testing, feedback loops)"
      - "Ship v2 ranking system"
      - "Set up offline benchmarks"
      - "Mentor new hires"
  If the JD does not describe concrete built-system deliverables, return an empty list — do NOT invent projects.
- position_requirement_type: "product" if the JD clearly names WHAT will be built / owned (specific systems, features, users, metrics to move) — even if the exact tech is left to the candidate. "profile" if the JD only lists required skills / years / titles / responsibilities without saying what concrete thing the candidate will produce. When ambiguous, prefer "profile".
- company_type: infer from the JD's self-description (funding stage, size, mention of "Series A", "MNC", "consulting", "product company", etc.). Use "startup" for early-stage (seed / Series A/B), "mid_stage" for Series C+/growth, "mnc" for large multinationals, "service" for IT-services / consulting shops (TCS, Infosys, Wipro, Accenture, Cognizant, Capgemini and similar), "agency" for creative/marketing shops, "other" if none fit.
- exclude_companies: list of specific companies the JD calls out as a NEGATIVE SIGNAL — the JD does not want candidates whose career has been (mostly) at these companies. This is a DOWN-WEIGHT list, not a hard-filter list — downstream ranking will penalize candidates proportionally to how long they've been at any of these companies. Extract each individual company name as its own list item, even if the JD lists them grouped in one sentence. Common pattern: JD names IT-services / consulting shops (TCS, Infosys, Wipro, Accenture, Cognizant, Capgemini) — list them individually. Do NOT include company names that were merely mentioned as examples of good career paths, and do NOT include Redrob or the hiring company itself. Return [] if the JD names no companies to down-weight.
- location.remote_ok: set true ONLY if the JD explicitly permits a fully-remote work arrangement (e.g. "100% remote", "work from anywhere", "no office attendance required"). The word "hybrid" ALONE does NOT mean remote — "hybrid" almost always means "some days in-office, some days WFH" and requires the candidate to live near the office, so remote_ok=false in that case. Similarly, "flexible cadence" / "flexible days" describes IN-OFFICE scheduling flexibility, not remote work — remote_ok=false. When the JD is silent or ambiguous, set remote_ok=false (safer default: assume onsite/hybrid).
- location.cities: cities where the role is BASED (offices, target hire locations). Do NOT include cities mentioned as "welcome to apply from" if they aren't office locations — those are still expected to relocate. Extract city names only, drop state/country.
"""


_FALLBACK_OUTPUT_YAML_TEMPLATE = """\
role_title:                # string — the JD's own role name
seniority:                 # one of: junior | mid | senior | staff | principal | null
min_years_experience:      # number | null
max_years_experience:      # number | null
required_technical_skills: # list of canonical short skill tokens
  - example_skill
preferred_technical_skills:
  - example_skill
target_role_titles:        # current titles that fit
  - example_title
domain_keywords:           # product/domain terms (RAG, ranking, embeddings, ...)
  - example_keyword
disqualifiers:             # explicit dealbreakers only
  - example_disqualifier
visa_sponsorship:          # bool; true only if JD explicitly offers it; default false
expected_projects:         # built systems, phrased as natural project descriptions for
                           # semantic search against candidates' career_history descriptions.
                           # 10-25 words each, engineer-voice, name the system + stack + problem domain.
  - "hybrid retrieval system combining BM25 and dense embeddings for candidate ranking"
  - "LLM-based re-ranker over top-k retrieval results with prompt-engineered relevance scoring"
position_requirement_type: # one of: product | profile
company_type:              # one of: startup | mid_stage | mnc | service | agency | other
exclude_companies:         # companies to DOWN-WEIGHT (not hard-reject) —
                           # candidates get a negative score proportional to
                           # months spent at any listed company.
  - example_company
location:
  country:                 # string — country the role is based in (e.g. 'India')
  cities:
    - example_city
  remote_ok:               # bool | null
  relocation_ok:           # bool | null
"""


# Load overrides from scripts/jd_compiler_prompt.md. Kept as a function so
# the auto-tuner can call reload_prompt() after rewriting the file mid-process.
def reload_prompt() -> tuple[str, str]:
    """Re-read jd_compiler_prompt.md and return (system_prompt, yaml_template).
    Falls back to the hardcoded strings when the file is missing or malformed."""
    loaded = _load_prompt_file(PROMPT_PATH)
    if loaded is None:
        print(
            f"[llm_JD_compiler] {PROMPT_PATH} missing or malformed — using fallback prompt.",
            file=sys.stderr,
        )
        return _FALLBACK_SYSTEM_PROMPT, _FALLBACK_OUTPUT_YAML_TEMPLATE
    return loaded


SYSTEM_PROMPT, OUTPUT_YAML_TEMPLATE = reload_prompt()


def build_user_prompt(jd_text: str) -> str:
    return (
        "Extract hiring signals from the JD below into a YAML document matching this shape "
        "(replace the placeholder values / comments with real values from the JD; keep the same keys and structure):\n\n"
        "```\n"
        f"{OUTPUT_YAML_TEMPLATE}"
        "```\n\n"
        "JOB DESCRIPTION:\n"
        "-----\n"
        f"{jd_text}\n"
        "-----\n\n"
        "Return only the YAML document. No prose, no code fences."
    )


def call_llm(jd_text: str, model: str, base_url: str, api_key: str, temperature: float, debug: bool = False) -> dict:
    try:
        from openai import OpenAI
    except ImportError as e:
        raise SystemExit("openai package not installed. Run: pip install openai") from e

    client = OpenAI(api_key=api_key, base_url=base_url)

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": build_user_prompt(jd_text)},
    ]

    # No response_format hint — we ask for YAML in the prompt. Most providers'
    # json_object mode either doesn't understand YAML or forces JSON output.
    resp = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
    )

    content = resp.choices[0].message.content or ""
    if debug:
        print(f"\n--- RAW LLM RESPONSE ({len(content)} chars) ---", file=sys.stderr)
        print(content[:4000] + ("...[truncated]" if len(content) > 4000 else ""), file=sys.stderr)
        print("--- END RAW RESPONSE ---\n", file=sys.stderr)
    if not content.strip():
        raise SystemExit("LLM returned empty content. Try --debug and inspect the raw response.")
    try:
        return _parse_yaml_lenient(content)
    except yaml.YAMLError as e:
        print("\n--- RAW LLM RESPONSE (unparseable) ---", file=sys.stderr)
        print(content, file=sys.stderr)
        print("--- END RAW RESPONSE ---\n", file=sys.stderr)
        raise SystemExit(f"Model output was not valid YAML: {e}")


def _parse_yaml_lenient(raw: str) -> dict:
    """Strip common wrappers (``` fences, leading prose) before yaml.safe_load.

    Handles the ways a chat model tends to violate 'YAML only':
      - ```yaml … ``` code fences
      - a chatty preamble before the first top-level key
      - a trailing 'Explanation: …' after the YAML block
    """
    s = raw.strip()

    # Strip a leading fence (```yaml or ``` or ```yml).
    if s.startswith("```"):
        first_nl = s.find("\n")
        if first_nl != -1:
            s = s[first_nl + 1 :]
        if s.rstrip().endswith("```"):
            s = s.rstrip()[:-3]
        s = s.strip()

    # If there's still a preamble, find the first line that looks like a YAML
    # top-level key and start from there.
    key_re = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*\s*:", re.MULTILINE)
    m = key_re.search(s)
    if m and m.start() > 0:
        s = s[m.start() :]

    data = yaml.safe_load(s)
    if not isinstance(data, dict):
        raise yaml.YAMLError(f"expected a YAML mapping at top level, got {type(data).__name__}")
    return data


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument(
        "--in",
        dest="inp",
        required=True,
        help="JD file: .docx, .txt, .md, or '-' for stdin",
    )
    ap.add_argument("--out", dest="out", required=True, help="Output JSON path")
    ap.add_argument(
        "--model",
        default=os.environ.get("LLM_MODEL_NAME") or os.environ.get("JD_MODEL"),
        help="Model id (defaults to $LLM_MODEL_NAME or $JD_MODEL)",
    )
    ap.add_argument(
        "--base-url",
        default=os.environ.get("LLM_BASE_URL") or os.environ.get("OPENAI_BASE_URL"),
        help="OpenAI-compatible base URL (defaults to $LLM_BASE_URL or $OPENAI_BASE_URL)",
    )
    ap.add_argument(
        "--api-key-env",
        default="LLM_API_KEY",
        help="Env var holding the API key (default LLM_API_KEY, falls back to OPENAI_API_KEY)",
    )
    ap.add_argument("--temperature", type=float, default=0.0)
    ap.add_argument("--debug", action="store_true", help="Print raw LLM response before parse")
    ap.add_argument(
        "--print-prompt",
        action="store_true",
        help="Print the compiled prompt and exit without calling the LLM",
    )
    args = ap.parse_args()

    jd_text = read_jd(Path(args.inp))

    if args.print_prompt:
        print(SYSTEM_PROMPT)
        print("\n---\n")
        print(build_user_prompt(jd_text))
        return 0

    api_key = os.environ.get(args.api_key_env) or os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise SystemExit(
            f"Missing API key: set ${args.api_key_env} or $OPENAI_API_KEY."
        )
    if not args.base_url:
        raise SystemExit("Missing base URL: set $LLM_BASE_URL or pass --base-url.")
    if not args.model:
        raise SystemExit("Missing model: set $LLM_MODEL_NAME or pass --model.")

    result = call_llm(
        jd_text=jd_text,
        model=args.model,
        base_url=args.base_url,
        api_key=api_key,
        temperature=args.temperature,
        debug=args.debug,
    )

    # Enforce documented defaults for fields the model may omit.
    if result.get("visa_sponsorship") is None:
        result["visa_sponsorship"] = False

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w") as f:
        json.dump(result, f, indent=2)

    req = result.get("required_technical_skills") or []
    pref = result.get("preferred_technical_skills") or []
    projects = result.get("expected_projects") or []
    print(f"Wrote {out_path}")
    print(f"  role: {result.get('role_title')}  seniority: {result.get('seniority')}")
    print(
        f"  yoe: {result.get('min_years_experience')} - {result.get('max_years_experience')}"
    )
    print(f"  company_type: {result.get('company_type')}  visa_sponsorship: {result.get('visa_sponsorship')}  position_requirement_type: {result.get('position_requirement_type')}")
    print(f"  required skills ({len(req)}): {', '.join(req[:8])}{'…' if len(req) > 8 else ''}")
    print(f"  preferred skills ({len(pref)}): {', '.join(pref[:8])}{'…' if len(pref) > 8 else ''}")
    print(f"  expected projects ({len(projects)}):")
    for pr in projects[:6]:
        print(f"    - {pr}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
